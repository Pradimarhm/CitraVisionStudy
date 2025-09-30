from docx import Document
from docx.shared import Pt

# Membuat dokumen baru
doc = Document()

# Judul dokumen
doc.add_heading('Penjelasan Kode Konvolusi dan Deteksi Tepi pada Citra Chelsea', level=1)

# Fungsi untuk menambahkan paragraf dengan ukuran font tertentu
def add_paragraph(text, size=11):
    p = doc.add_paragraph(text)
    run = p.runs[0]
    run.font.size = Pt(size)

# Menambahkan isi penjelasan
add_paragraph("Berikut penjelasan rinci dari kode yang diberikan, bagian per bagian:", 12)

# 1. Import Library
doc.add_heading('1. Import Library', level=2)
add_paragraph(
    "Kode mengimpor beberapa library penting:\n"
    "- numpy: untuk operasi array dan matriks.\n"
    "- matplotlib.pyplot: untuk menampilkan gambar.\n"
    "- skimage.data: untuk mengambil contoh citra chelsea.\n"
    "- skimage.color: untuk konversi citra ke grayscale.\n"
    "- scipy.ndimage.convolve: untuk konvolusi otomatis dengan padding refleksi.\n"
    "- cv2 (OpenCV): untuk deteksi tepi Canny."
)

# 2. Load dan Preprocessing Citra
doc.add_heading('2. Load dan Preprocessing Citra', level=2)
add_paragraph(
    "Citra chelsea diambil dari skimage dan diubah ke grayscale:\n"
    "- data.chelsea() mengambil citra kucing berwarna.\n"
    "- color.rgb2gray() mengubah citra RGB menjadi grayscale dengan nilai intensitas 0-1."
)

# 3. Definisi Kernel 5x5 untuk Konvolusi
doc.add_heading('3. Definisi Kernel 5x5 untuk Konvolusi', level=2)
add_paragraph(
    "Kernel Gaussian blur 5x5 didefinisikan dan dinormalisasi agar total bobot kernel menjadi 1.\n"
    "Hal ini menjaga intensitas citra agar tidak berubah secara signifikan setelah konvolusi."
)

# 4. Konvolusi Manual dengan Padding Refleksi
doc.add_heading('4. Konvolusi Manual dengan Padding Refleksi', level=2)
add_paragraph(
    "Konvolusi manual dilakukan dengan langkah berikut:\n"
    "- Padding citra dengan mode 'reflect' menggunakan np.pad.\n"
    "- Melakukan iterasi pada setiap piksel citra asli.\n"
    "- Mengambil region 5x5 dari citra yang sudah dipadding.\n"
    "- Mengalikan region dengan kernel dan menjumlahkan hasilnya.\n"
    "- Menyimpan hasil pada citra output."
)

# 5. Konvolusi Otomatis dengan scipy.ndimage.convolve
doc.add_heading('5. Konvolusi Otomatis dengan scipy.ndimage.convolve', level=2)
add_paragraph(
    "Konvolusi otomatis menggunakan fungsi convolve dari scipy dengan mode padding 'reflect'.\n"
    "Metode ini lebih efisien dan menghasilkan output yang sama dengan konvolusi manual."
)

# 6. Deteksi Tepi Sobel dengan Kernel 5x5
doc.add_heading('6. Deteksi Tepi Sobel dengan Kernel 5x5', level=2)
add_paragraph(
    "Kernel Sobel 3x3 diperbesar menjadi 5x5 untuk mendeteksi perubahan intensitas horizontal dan vertikal.\n"
    "Gradien arah x dan y dihitung dengan konvolusi, lalu magnitude gradien dihitung dengan np.hypot dan dinormalisasi."
)

# 7. Deteksi Tepi Canny
doc.add_heading('7. Deteksi Tepi Canny', level=2)
add_paragraph(
    "Deteksi tepi Canny menggunakan OpenCV:\n"
    "- Citra grayscale dikonversi ke uint8 (0-255).\n"
    "- Fungsi cv2.Canny digunakan dengan threshold rendah 150 dan tinggi 250.\n"
    "- Hasil berupa citra biner tepi."
)

# 8. Deteksi Tepi Prewitt dengan Kernel 5x5
doc.add_heading('8. Deteksi Tepi Prewitt dengan Kernel 5x5', level=2)
add_paragraph(
    "Kernel Prewitt 3x3 diperbesar menjadi 5x5 untuk mendeteksi perubahan intensitas.\n"
    "Gradien arah x dan y dihitung dan magnitude gradien dinormalisasi."
)

# 9. Visualisasi Hasil
doc.add_heading('9. Visualisasi Hasil', level=2)
add_paragraph(
    "Hasil ditampilkan dalam grid 2x3:\n"
    "- Citra asli dan hasil konvolusi dengan colormap grayscale.\n"
    "- Hasil deteksi tepi Sobel, Canny, dan Prewitt dengan colormap berbeda (inferno, plasma, viridis) agar mudah dibedakan.\n"
    "- Sumbu gambar dihilangkan untuk tampilan lebih bersih."
)

# Kesimpulan
doc.add_heading('Kesimpulan', level=2)
add_paragraph(
    "- Konvolusi manual mengajarkan cara kerja dasar konvolusi dengan padding refleksi.\n"
    "- Konvolusi otomatis lebih efisien dan praktis.\n"
    "- Deteksi tepi Sobel dan Prewitt menggunakan kernel 5x5 yang diperbesar.\n"
    "- Deteksi tepi Canny menggunakan threshold ganda untuk hasil presisi.\n"
    "- Penggunaan colormap berbeda membantu membedakan hasil deteksi tepi secara visual."
)

# Simpan dokumen
doc.save('Penjelasan_Kode_Konvolusi_Dan_Deteksi_Tepi.docx')

print("Dokumen berhasil dibuat dengan nama 'Penjelasan_Kode_Konvolusi_Dan_Deteksi_Tepi.docx'")
